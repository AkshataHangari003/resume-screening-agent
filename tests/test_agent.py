import sys
from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas


# Allow imports from app/
BASE_DIR = Path(__file__).resolve().parent.parent
APP_DIR = BASE_DIR / "app"

sys.path.insert(0, str(APP_DIR))


from parser import extract_resume_text

from extractor import (
    extract_name,
    extract_skills,
    extract_education,
    extract_experience,
    extract_required_skills,
)

from scorer import (
    calculate_skill_match,
    calculate_experience_score,
    extract_experience_years,
)

from ranker import (
    get_matched_skills,
    get_missing_skills,
    generate_recommendation,
)


def test_txt_parsing(tmp_path):
    """Test TXT resume parsing."""

    file_path = tmp_path / "resume.txt"

    file_path.write_text(
        """NAME: Test Candidate

EDUCATION:
MCA

EXPERIENCE:
1 year software development experience

SKILLS:
Python, Java, SQL, Git
""",
        encoding="utf-8",
    )

    text = extract_resume_text(str(file_path))

    assert "Test Candidate" in text
    assert "Python" in text


def test_pdf_parsing(tmp_path):
    """Test PDF resume parsing."""

    file_path = tmp_path / "resume.pdf"

    pdf = canvas.Canvas(str(file_path))

    pdf.drawString(
        50,
        800,
        "NAME: PDF Candidate",
    )

    pdf.drawString(
        50,
        770,
        "SKILLS: Python, SQL, Git",
    )

    pdf.save()

    text = extract_resume_text(str(file_path))

    assert "PDF Candidate" in text
    assert "Python" in text


def test_docx_parsing(tmp_path):
    """Test DOCX resume parsing."""

    file_path = tmp_path / "resume.docx"

    document = Document()

    document.add_paragraph(
        "NAME: DOCX Candidate"
    )

    document.add_paragraph(
        "SKILLS: Python, Java, SQL"
    )

    document.save(str(file_path))

    text = extract_resume_text(str(file_path))

    assert "DOCX Candidate" in text
    assert "Python" in text


def test_name_extraction():
    """Test candidate name extraction."""

    text = "NAME: Rahul Kumar"

    assert extract_name(text) == "Rahul Kumar"


def test_skill_extraction():
    """Test candidate skill extraction."""

    text = """
SKILLS:
Python, Java, SQL, Git, GitHub
"""

    skills = extract_skills(text)

    assert "Python" in skills
    assert "Java" in skills
    assert "SQL" in skills
    assert "Git" in skills
    assert "GitHub" in skills


def test_education_extraction():
    """Test education extraction."""

    text = """
EDUCATION:
Master of Computer Applications

EXPERIENCE:
1 year experience
"""

    education = extract_education(text)

    assert education == "Master of Computer Applications"


def test_experience_extraction():
    """Test experience extraction."""

    text = """
EXPERIENCE:
1 year software development experience
"""

    experience = extract_experience(text)

    assert experience == "1 year software development experience"


def test_required_skill_extraction():
    """Test required skill extraction from JD."""

    jd = """
JOB TITLE:
Junior Software Engineer

REQUIRED SKILLS:
- Python
- Java
- SQL
- Git
- GitHub
- REST APIs
- Data Structures
- Algorithms
- Object-Oriented Programming

PREFERRED SKILLS:
- Docker
"""

    skills = extract_required_skills(jd)

    assert "Python" in skills
    assert "Java" in skills
    assert "SQL" in skills
    assert "Git" in skills
    assert "GitHub" in skills
    assert "REST APIs" in skills
    assert "Data Structures" in skills
    assert "Algorithms" in skills
    assert "Object-Oriented Programming" in skills


def test_git_and_github_are_separate():
    """Ensure Git is not incorrectly detected inside GitHub."""

    jd = """
REQUIRED SKILLS:
- GitHub

PREFERRED SKILLS:
- Docker
"""

    skills = extract_required_skills(jd)

    assert "GitHub" in skills
    assert "Git" not in skills


def test_skill_match():
    """Test skill matching percentage."""

    required = [
        "Python",
        "Java",
        "SQL",
        "Git",
    ]

    candidate = [
        "Python",
        "Java",
    ]

    score = calculate_skill_match(
        required,
        candidate,
    )

    assert score == 50.0


def test_experience_years():
    """Test experience extraction."""

    assert (
        extract_experience_years(
            "6 months software development internship"
        )
        == 0.5
    )

    assert (
        extract_experience_years(
            "2 years software engineer experience"
        )
        == 2.0
    )

    assert (
        extract_experience_years(
            "Fresh graduate"
        )
        == 0.0
    )


def test_experience_score():
    """Test experience scoring."""

    assert (
        calculate_experience_score(
            "1 year software development experience"
        )
        == 100.0
    )

    assert (
        calculate_experience_score(
            "6 months internship"
        )
        == 85.0
    )

    assert (
        calculate_experience_score(
            "Fresh graduate"
        )
        == 70.0
    )


def test_matched_skills():
    """Test matched required skills."""

    required = [
        "Python",
        "Java",
        "SQL",
    ]

    candidate = [
        "Python",
        "SQL",
    ]

    matched = get_matched_skills(
        required,
        candidate,
    )

    assert matched == [
        "Python",
        "SQL",
    ]


def test_missing_skills():
    """Test missing required skills."""

    required = [
        "Python",
        "Java",
        "SQL",
    ]

    candidate = [
        "Python",
    ]

    missing = get_missing_skills(
        required,
        candidate,
    )

    assert missing == [
        "Java",
        "SQL",
    ]


def test_recommendations():
    """Test recommendation thresholds."""

    assert (
        generate_recommendation(85)
        == "Strong Match"
    )

    assert (
        generate_recommendation(70)
        == "Good Match"
    )

    assert (
        generate_recommendation(55)
        == "Moderate Match"
    )

    assert (
        generate_recommendation(40)
        == "Weak Match"
    )